"""Render reports to Markdown, Word (.docx), PDF and plain-text transcripts."""

from __future__ import annotations

import logging
import os
import re
from typing import List

from .models import TranscriptSegment, VideoInfo, VideoResult
from .summarize import _fmt_ts

log = logging.getLogger("usum.render")


# ----------------------------------------------------------------------------- #
# Building the combined Markdown report
# ----------------------------------------------------------------------------- #
def _duration_str(seconds) -> str:
    if not seconds:
        return "unknown"
    return _fmt_ts(seconds)


def _date_str(yyyymmdd) -> str:
    if not yyyymmdd or len(str(yyyymmdd)) != 8:
        return "unknown"
    s = str(yyyymmdd)
    return f"{s[0:4]}-{s[4:6]}-{s[6:8]}"


def _meta_block(r: VideoResult) -> str:
    info = r.info
    return (
        f"- **Channel:** {info.uploader or 'unknown'}  \n"
        f"- **Duration:** {_duration_str(info.duration)}  \n"
        f"- **Published:** {_date_str(info.upload_date)}  \n"
        f"- **URL:** {info.url}  \n"
        f"- **Transcript source:** {r.transcript_source}"
    )


def build_video_markdown(r: VideoResult, heading_level: int = 2) -> str:
    """Markdown for a single video: title heading + metadata + summary."""
    h = "#" * heading_level
    lines = [f"{h} {r.info.title}", "", _meta_block(r), ""]
    if r.error:
        lines.append(f"> ⚠️ Could not summarise this video: {r.error}")
    else:
        lines.append(r.summary_markdown.strip())
    return "\n".join(lines).rstrip() + "\n"


def build_markdown(results: List[VideoResult], generated_on: str) -> str:
    lines: List[str] = []
    lines.append("# uSum — YouTube Summary Report")
    lines.append("")
    lines.append(f"*Generated {generated_on} · {len(results)} video(s)*")
    lines.append("")
    if len(results) > 1:
        lines.append("## Contents")
        for i, r in enumerate(results, 1):
            lines.append(f"{i}. {r.info.title}")
        lines.append("")

    for r in results:
        lines.append(build_video_markdown(r, heading_level=2))
        lines.append("---")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def write_markdown(markdown: str, path: str) -> None:
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(markdown)
    log.info("Wrote %s", path)


def write_transcript(info: VideoInfo, segments: List[TranscriptSegment], path: str) -> None:
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(f"{info.title}\n{info.url}\n\n")
        for s in segments:
            fh.write(f"[{_fmt_ts(s.start)}] {s.text}\n")
    log.info("Wrote %s", path)


# ----------------------------------------------------------------------------- #
# Word (.docx) — small Markdown subset parser
# ----------------------------------------------------------------------------- #
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_runs(paragraph, text: str) -> None:
    """Add text to a docx paragraph, honouring **bold** spans."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def write_docx(markdown: str, path: str) -> None:
    from docx import Document

    doc = Document()
    for raw in markdown.split("\n"):
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped in ("---", "***", "___"):
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(heading.group(2), level=level)
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            _add_runs(doc.add_paragraph(style="List Bullet"), bullet.group(1))
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            _add_runs(doc.add_paragraph(style="List Number"), numbered.group(1))
            continue
        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs(p, quote.group(1))
            continue
        _add_runs(doc.add_paragraph(), stripped)

    doc.save(path)
    log.info("Wrote %s", path)


# ----------------------------------------------------------------------------- #
# PDF — Markdown -> HTML -> fpdf2, with a Unicode font when one can be found
# ----------------------------------------------------------------------------- #
# Each entry maps fpdf2 style codes ("", B, I, BI) to a TTF. The "" (regular)
# file must exist for the set to be used; missing variants fall back to regular,
# so write_html never hits an undefined font when it renders bold/italic spans.
_FONT_SETS = [
    {  # Windows Arial
        "": r"C:\Windows\Fonts\arial.ttf",
        "B": r"C:\Windows\Fonts\arialbd.ttf",
        "I": r"C:\Windows\Fonts\ariali.ttf",
        "BI": r"C:\Windows\Fonts\arialbi.ttf",
    },
    {  # Linux DejaVu
        "": "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "B": "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "I": "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
        "BI": "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf",
    },
    {  # Linux Liberation
        "": "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "B": "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "I": "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf",
        "BI": "/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf",
    },
    {  # macOS
        "": "/Library/Fonts/Arial.ttf",
        "B": "/Library/Fonts/Arial Bold.ttf",
        "I": "/Library/Fonts/Arial Italic.ttf",
        "BI": "/Library/Fonts/Arial Bold Italic.ttf",
    },
]


def _register_unicode_font(pdf, family: str = "body") -> bool:
    """Register a Unicode font (all 4 styles) on the PDF. Returns True on success."""
    for fonts in _FONT_SETS:
        regular = fonts.get("")
        if not regular or not os.path.exists(regular):
            continue
        for style in ("", "B", "I", "BI"):
            path = fonts.get(style) or regular
            if not os.path.exists(path):
                path = regular
            pdf.add_font(family, style, path)
        return True
    return False


def write_pdf(markdown: str, path: str) -> None:
    import markdown as md
    from fpdf import FPDF

    html = md.markdown(markdown, extensions=["extra"])

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    if _register_unicode_font(pdf):
        pdf.set_font("body", size=11)
    else:
        # Core font only handles latin-1; transliterate anything outside it.
        log.warning("No Unicode TTF found; PDF will use a latin-1 core font.")
        pdf.set_font("helvetica", size=11)
        html = html.encode("latin-1", "replace").decode("latin-1")

    pdf.write_html(html)
    pdf.output(path)
    log.info("Wrote %s", path)


# ----------------------------------------------------------------------------- #
# Dispatch
# ----------------------------------------------------------------------------- #
def render_outputs(
    results: List[VideoResult],
    out_dir: str,
    formats: List[str],
    generated_on: str,
    basename: str,
) -> List[str]:
    os.makedirs(out_dir, exist_ok=True)
    written: List[str] = []

    markdown = build_markdown(results, generated_on)

    if "md" in formats:
        p = os.path.join(out_dir, f"{basename}.md")
        write_markdown(markdown, p)
        written.append(p)

    if "docx" in formats:
        p = os.path.join(out_dir, f"{basename}.docx")
        try:
            write_docx(markdown, p)
            written.append(p)
        except Exception as exc:
            log.error("Failed to write .docx: %s", exc)

    if "pdf" in formats:
        p = os.path.join(out_dir, f"{basename}.pdf")
        try:
            write_pdf(markdown, p)
            written.append(p)
        except Exception as exc:
            log.error("Failed to write .pdf: %s", exc)

    if "txt" in formats:
        tdir = os.path.join(out_dir, "transcripts")
        os.makedirs(tdir, exist_ok=True)
        for r in results:
            if not r.segments:
                continue
            safe = re.sub(r"[^\w\-]+", "_", r.info.title)[:80] or r.info.id
            p = os.path.join(tdir, f"{safe}.txt")
            write_transcript(r.info, r.segments, p)
            written.append(p)

    return written
